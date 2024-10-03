import docx  # Import the 'docx' library to create and manipulate Word documents
import pandas as pd  # Import the 'pandas' library to handle Excel data

# Load the Excel file using a raw string (the 'r' in front of the path helps avoid issues with backslashes in file paths)
excel_file = r'C:\Users\akshay.a\Desktop\Test Cases Artifact Automation\LinkedIn_Test_Cases.xlsx'
data = pd.read_excel(excel_file)  # Read the Excel file into a pandas DataFrame

# Create a dictionary to store rows of data grouped by the 'ID' column
id_data = {}  # Initialize an empty dictionary
for index, row in data.iterrows():  # Loop through each row in the DataFrame
    if row['ID'] not in id_data:  # Check if this 'ID' is not already in the dictionary
        id_data[row['ID']] = []  # If not, create an empty list for that 'ID'
    id_data[row['ID']].append(row)  # Append the current row to the list of rows for that 'ID'

# For each unique 'ID', create a separate Word document
for id, rows in id_data.items():
    doc = docx.Document()  # Create a new Word document
    
    # Add the title (from the first row of this group of rows)
    title = rows[0]['Title']  # Get the title from the first row under 'Title' column
    doc.add_heading(str(title), 0)  # Add the title as a heading at level 0 (largest size)
    
    # Add a blank line for better formatting
    doc.add_paragraph("")
    
    # Add a line indicating who created the document
    doc.add_paragraph("Created by: Akshay A Palande")
    
    # Add another blank line for spacing
    doc.add_paragraph("")
    
    # Add a table to display the test steps from the Excel file
    # The table will have as many rows as there are steps + 1 header row, and 3 columns (Step, Action, Expected)
    table = doc.add_table(rows=len(rows)+1, cols=3, style='Table Grid')
    
    # Define the header for the table (first row with titles)
    table.cell(0, 0).text = "Test Step"  # First column header
    table.cell(0, 1).text = "Step Action"  # Second column header
    table.cell(0, 2).text = "Step Expected"  # Third column header
    
    # Extract the test steps data from the Excel file
    test_steps = []  # Create an empty list to store test steps
    for index, row in enumerate(rows):  # Loop through each row for the current 'ID'
        if 'Test Step' in row and pd.notna(row['Test Step']):  # Check if 'Test Step' exists and is not empty
            test_step = row['Test Step']  # Get the test step
            step_action = row['Step Action'] if 'Step Action' in row and pd.notna(row['Step Action']) else ''  # Get the step action if it exists
            step_expected = row['Step Expected'] if 'Step Expected' in row and pd.notna(row['Step Expected']) else ''  # Get the expected result if it exists
            test_steps.append((test_step, step_action, step_expected))  # Append this data as a tuple to the list
    
    # Fill in the test steps table with the data from the list
    for i, (test_step, step_action, step_expected) in enumerate(test_steps):  # Loop through each test step
        table.cell(i+1, 0).text = str(test_step)  # Add the test step to the first column
        table.cell(i+1, 1).text = str(step_action)  # Add the step action to the second column
        table.cell(i+1, 2).text = str(step_expected)  # Add the expected result to the third column
    
    # Add another blank line after the table for better spacing
    doc.add_paragraph("")
    
    # Add a section for 'Test Data' information
    doc.add_paragraph("Test Data:")  # Add the 'Test Data' title
    doc.add_paragraph(str(rows[0]['Test Data']))  # Add the actual test data (assuming it's in the first row for this 'ID')
    
    # Add another blank line for spacing
    doc.add_paragraph("")
    
    # Add the title again at the end (optional, can be removed if not needed)
    doc.add_paragraph("")  # Add a blank line for spacing
    doc.add_heading(str(title), 0)  # Add the title as a heading again
    
    # Save the Word document using the 'Title' as the filename
    doc.save(f"{str(title)}.docx")  # Save the document with the title as the file name
